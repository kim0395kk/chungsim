# -*- coding: utf-8 -*-
"""
Govable AI - DOCX 문서 생성 모듈
python-docx를 사용한 공문서 및 보고서 생성
"""
import io
from typing import Dict, Any, List, Optional
from datetime import datetime

try:
    from docx import Document
    from docx.shared import Pt, RGBColor, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False
    Document = None


def generate_official_docx(doc_data: Dict[str, Any]) -> bytes:
    """
    공문서 DOCX 생성
    
    Args:
        doc_data: 문서 데이터
            - title: 문서 제목
            - receiver: 수신자 (선택)
            - body_paragraphs: 본문 문단 리스트
            - department_head: 부서장 이름 (선택)
            - doc_num: 문서번호 (선택)
            
    Returns:
        DOCX 파일 바이트
    """
    if not DOCX_AVAILABLE:
        raise ImportError("python-docx 패키지가 설치되지 않았습니다. pip install python-docx")
    
    doc = Document()
    
    # 문서 제목
    title = doc_data.get('title', '공문서')
    heading = doc.add_heading(title, level=1)
    heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # 문서 정보
    doc_num = doc_data.get('doc_num', '')
    today = datetime.now().strftime('%Y년 %m월 %d일')
    
    if doc_num:
        p = doc.add_paragraph()
        p.add_run(f"문서번호: {doc_num}").bold = True
    
    p = doc.add_paragraph()
    p.add_run(f"시행일자: {today}").bold = True
    
    receiver = doc_data.get('receiver', '')
    if receiver:
        p = doc.add_paragraph()
        p.add_run(f"수신: {receiver}").bold = True
    
    doc.add_paragraph()  # 빈 줄
    
    # 본문
    body_paragraphs = doc_data.get('body_paragraphs', [])
    if isinstance(body_paragraphs, str):
        body_paragraphs = body_paragraphs.split('\n\n')
    
    for para_text in body_paragraphs:
        if para_text.strip():
            p = doc.add_paragraph(para_text.strip())
            p.style = 'Normal'
    
    doc.add_paragraph()  # 빈 줄
    
    # 발신 정보
    department_head = doc_data.get('department_head', '')
    if department_head:
        p = doc.add_paragraph()
        p.add_run(f"{department_head}").bold = True
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    
    # 바이트로 변환
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer.getvalue()


def generate_guide_docx(guide_data: Dict[str, Any]) -> bytes:
    """
    처리가이드 보고서 DOCX 생성
    
    Args:
        guide_data: 가이드 데이터
            - analysis: 분석 결과 딕셔너리
            - summary: 요약 (선택)
            - timeline: 처리 절차 (선택)
            - legal_basis: 법적 근거 (선택)
            
    Returns:
        DOCX 파일 바이트
    """
    if not DOCX_AVAILABLE:
        raise ImportError("python-docx 패키지가 설치되지 않았습니다. pip install python-docx")
    
    doc = Document()
    
    # 제목
    title = guide_data.get('title', '민원 처리 가이드')
    heading = doc.add_heading(title, level=1)
    heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # 작성일
    today = datetime.now().strftime('%Y년 %m월 %d일')
    p = doc.add_paragraph()
    p.add_run(f"작성일: {today}").italic = True
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    
    doc.add_paragraph()  # 빈 줄
    
    # 분석 결과
    analysis = guide_data.get('analysis', {})
    
    # 요약
    if 'summary' in analysis or 'summary' in guide_data:
        doc.add_heading('📋 요약', level=2)
        summary = analysis.get('summary', guide_data.get('summary', ''))
        if summary:
            doc.add_paragraph(str(summary))
        doc.add_paragraph()
    
    # 처리 절차
    if 'timeline' in analysis or 'timeline' in guide_data:
        doc.add_heading('📅 처리 절차', level=2)
        timeline = analysis.get('timeline', guide_data.get('timeline', []))
        
        if isinstance(timeline, list):
            for i, step in enumerate(timeline, 1):
                if isinstance(step, dict):
                    step_text = step.get('description', str(step))
                else:
                    step_text = str(step)
                doc.add_paragraph(f"{i}. {step_text}", style='List Number')
        else:
            doc.add_paragraph(str(timeline))
        doc.add_paragraph()
    
    # 법적 근거
    if 'legal_basis' in analysis or 'legal_basis' in guide_data:
        doc.add_heading('⚖️ 법적 근거', level=2)
        legal = analysis.get('legal_basis', guide_data.get('legal_basis', ''))
        if legal:
            doc.add_paragraph(str(legal))
        doc.add_paragraph()
    
    # 처리 전략
    if 'strategy' in analysis or 'strategy' in guide_data:
        doc.add_heading('🎯 처리 전략', level=2)
        strategy = analysis.get('strategy', guide_data.get('strategy', ''))
        if strategy:
            doc.add_paragraph(str(strategy))
        doc.add_paragraph()
    
    # 기타 필드들 (analysis 딕셔너리의 나머지)
    known_fields = {'summary', 'timeline', 'legal_basis', 'strategy', 'title'}
    for key, value in analysis.items():
        if key not in known_fields and value:
            # key를 문자열로 변환 후 처리
            key_str = str(key).replace("_", " ").title()
            doc.add_heading(f'📌 {key_str}', level=2)
            if isinstance(value, list):
                for item in value:
                    doc.add_paragraph(f"• {item}", style='List Bullet')
            else:
                doc.add_paragraph(str(value))
            doc.add_paragraph()
    
    # 바이트로 변환
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer.getvalue()
